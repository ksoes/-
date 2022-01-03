from flask import Flask, render_template, request, redirect, url_for
from werkzeug.utils import secure_filename
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import olefile
import docx 
import re
from konlpy.tag import Mecab
import pandas as pd
from tqdm import tqdm
from sklearn.feature_extraction.text import CountVectorizer
import sqlalchemy
from sqlalchemy import create_engine
import pymysql
import requests
from pandas.io.json import json_normalize
import os
import numpy as np
from IPython import display
from newspaper import Article
from sklearn.feature_extraction.text import TfidfVectorizer # tf-idf
from sklearn.metrics.pairwise import linear_kernel # tf-idf

# Flask 객체 인스턴스 생성
app = Flask(__name__) # __name__은 현재 실행중인 모듈이름을 전달

def connect_sql() :
    global engine
    global conn 
    engine = create_engine("mysql+pymysql://아이디:비밀번호@호스트주소:포트번호/DB이름?charset=utf8", encoding='utf-8') # mysql 주소
    conn = engine.connect() # mysql 연결
    return engine, conn

# --------------------------------------------------------

# 각 형식별 파일에서 텍스트 추출
def convert_pdf_to_text(url) :
    # pdf파일 읽어서 기사 추출
    rsrcmgr = PDFResourceManager()
    retstr=StringIO()
    codec='utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = open(url, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password=""
    maxpages=0
    caching=True
    pagenos=set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching, check_extractable=True) :
        interpreter.process_page(page)
    text = retstr.getvalue()
    fp.close()
    device.close()
    retstr.close()
    return text

def convert_hwp_to_text(url) :
    # olefile로 한글파일 열기
    f = olefile.OleFileIO(url)
    # PrvText 스트림 안의 내용 꺼내기 (유니코드 인코딩 되어있음)
    encoded_text = f.openstream('PrvText').read()
    # 유니코드이므로 UTF-16으로 디코딩
    decoded_text = encoded_text.decode('UTF-16')
    return decoded_text

def convert_txt_to_text(url) :
    f = open(url)
    # 파일의 내용 전체를 문자열로 리턴
    text = f.read()
    v = text
    f.close()
    return v

def convert_docx_to_text(url) :
    # .docx형식 파일 읽기
    s=[]
    dox = docx.Document(url)
    for i in dox.paragraphs :
        # 한 줄씩 읽어와서 리스트 s에 추가
        s.append(i.text)
    # join메소드로 리스트s의 원소들을 공백을 이용해서 구분
    text = ' '.join(s)
    return text

def convert_link_to_text(url) :
  # url읽어와서 본문 가져오기
  article = Article(url, language='ko')
  article.download()
  article.parse()
  text = article.text
  return text

def clean_sentence(v) :
    # 상단에 쓸모없는 단어 제거
    v = re.sub('인쇄하기|창닫기|인쇄|재배포|기자|무단전재|저작권자', '', v)
    # 바이트변수 제거
    if '\x00' in v :
        v = ' '.join(v.split('\x00'))
    if '\xa0' in v :
        v = ''.join(v.split('\xa0'))
    if '\x0c' in v :
        v = ''.join(v.split('\x0c'))
    # \n, 띄어쓰기 제거
    v = v.replace("\n","")
    v = v.replace(" ","")
    return v

def make_coword() :
    # 동시출현단어쌍 만들기 <일반기사, 가짜기사>
    # title_list, stop_word, 기사 본문
    # DTM.csv 파일 생성
    
    # 타이틀 리스트를 불러와서 title_list 변수에 저장한다.
    t_file_name = open('/home/ubuntu/etc_file/title_list.txt', 'r', encoding='utf-8')
    title_list = []
    for line in t_file_name.readlines():
        # txt파일을 readlines로 불러오면 개행 문자도 함께 읽어오기 때문에 인덱싱으로 처리해준다.
        title_list.append(line[:])
    t_file_name.close()
    
    
    #if s_file_name in locals :
    # 불용어 파일
    s_file_name = open('/home/ubuntu/web-codding/etc_file/stop_word.txt', 'r', encoding='utf-8') 
    stop_words_list = []
    for line in s_file_name.readlines() :
        stop_words_list.append(line.rstrip())
    s_file_name.close()
        
    
    # pandas의 read_csv 함수를 이용하여 csv 파일을 불러온다.
    # dataframe누적시킨걸로 csv저장한 파일 -> 추출된 본문 기사.csv 
    dataset = pd.read_csv('/home/ubuntu/web-codding/etc_file/organized_total_body.csv')
    
    # 각 형태소별로 분류(Tagging)해주는 Okt 객체를 불러온다.
    mecab = Mecab()
    
    for title in tqdm(title_list, desc='타이틀 리스트 진행도'):  # title_list에 대해 반복문을 실행
        # 각 타이틀에 대한 6770개 문서의 DTM을 표현하기 위해
        # CountVectorizer 객체를 선언
        cv = CountVectorizer()

        # 각 문서들의 말뭉치(corpus)를 저장할 리스트 선언
        corpus = []
        
        # 각 타이틀에 대한 문서들의 말 뭉치를 저장한다. (데이터가 많으면 이 부분에서 장시간이 소요될 수 있다.)
        for doc_num in range(len(dataset)) :
            # 각 말뭉치에서 명사 리스트를 만든다.
            noun_list = mecab.nouns(dataset[title].loc[doc_num])
            
            # 이를 문자열로 저장해야하기 때문에 join함수로 공백으로 구분해 corpus에 append한다.
            corpus.append(' '.join(noun_list))
        
        # CountVectorizer의 fit_transform 함수를 통해 DTM을 한번에 생성할 수 있다.
        DTM_Array = cv.fit_transform(corpus).toarray()
        
        # feature_names 함수를 사용하면 DTM의 각 열(column)이 어떤 단어에 해당하는지 알 수 있다.
        feature_names = cv.get_feature_names()
        
        # 추출해낸 데이터를 DataFrame 형식으로 변환한다.
        DTM_DataFrame = pd.DataFrame(DTM_Array, columns=feature_names)
            
            
        #if s_file_name in locals :
        # 불용어 열 제거
        DTM_DataFrame.drop(stop_words_list, axis='columns', inplace=True)
            
        
        # 최종적으로 DTM을 csv 파일로 저장한다.
        DTM_DataFrame.to_csv('/home/ubuntu/web-codding/etc_file/_trustnews_DTM.csv', encoding='utf-8-sig')
            
        
        # DTM파일 가지고 동시출현단어.csv 파일 만들기
        dataset = pd.read_csv('/home/ubuntu/web-codding/etc_file/_trustnews_DTM.csv') 

        column_list = dataset.columns[1:]
        word_length = len(column_list)
        
        count_dict={}
        
        for doc_number in tqdm(range(len(dataset)), desc='단어쌍 만들기 진행중') :
            tmp = dataset.loc[doc_number]
        for i, word1 in enumerate(column_list) :
            if tmp[word1] :
                for j in range(i+1, word_length) :
                    if tmp[column_list[j]] :
                        count_dict[column_list[i], column_list[j]] = count_dict.get((column_list[i]), 0) + max(tmp[word1], tmp[column_list[j]])
        
        count_list=[]
        for words in count_dict :
            count_list.append([words[0], words[1], count_dict[words]])
        
        df = pd.DataFrame(count_list, columns=["word1", "word2", "freq"])
        df = df.sort_values(by=['freq'], ascending=False) # 내림차순
        df = df.reset_index(drop=True)
         
        df.to_csv('/home/ubuntu/web-codding/etc_file/_trustnews_networkx.csv', encoding='utf-8-sig') 



def make_fcoword() :
    # 가짜뉴스_동시출현빈도 파일 생성
    
    # 타이틀 리스트를 불러와서 title_list 변수에 저장한다.
    t_file_name = open('/home/ubuntu/web-codding/etc_file/title_list.txt', 'r', encoding='utf-8') 
    title_list = []
    for line in t_file_name.readlines():
        # txt파일을 readlines로 불러오면 개행 문자도 함께 읽어오기 때문에 인덱싱으로 처리해준다.
        title_list.append(line[:])
    t_file_name.close()
    
    """ 불용어 제거(일치하는게 없으면 오류남)  
    #if s_file_name in locals :
    # 불용어 파일
    s_file_name = open('./stop_word.txt', 'r', encoding='utf-8')
    stop_words_list = []
    for line in s_file_name.readlines() :
        stop_words_list.append(line.rstrip())
    s_file_name.close()
    """   

    # pandas의 read_csv 함수를 이용하여 csv 파일을 불러온다.
    # dataframe누적시킨걸로 csv저장한 파일 -> 추출된 본문 기사.csv
    dataset = pd.read_csv('/home/ubuntu/web-codding/etc_file/fakenewsbody_sum.csv')
        
    # 각 형태소별로 분류(Tagging)해주는 Okt 객체를 불러온다.
    mecab = Mecab()
    
    for title in tqdm(title_list, desc='타이틀 리스트 진행도'):  # title_list에 대해 반복문을 실행
        # 각 타이틀에 대한 6770개 문서의 DTM을 표현하기 위해
        # CountVectorizer 객체를 선언
        cv = CountVectorizer()
        
        # 각 문서들의 말뭉치(corpus)를 저장할 리스트 선언
        corpus = []
        
        # 각 타이틀에 대한 문서들의 말 뭉치를 저장한다. (데이터가 많으면 이 부분에서 장시간이 소요될 수 있다.)
        for doc_num in range(len(dataset)) :
            # 각 말뭉치에서 명사 리스트를 만든다.
            noun_list = mecab.nouns(dataset[title].loc[doc_num])
            
            # 이를 문자열로 저장해야하기 때문에 join함수로 공백으로 구분해 corpus에 append한다.
            corpus.append(' '.join(noun_list))
        
        # CountVectorizer의 fit_transform 함수를 통해 DTM을 한번에 생성할 수 있다.
        DTM_Array = cv.fit_transform(corpus).toarray()
        
        # feature_names 함수를 사용하면 DTM의 각 열(column)이 어떤 단어에 해당하는지 알 수 있다.
        feature_names = cv.get_feature_names()
        
        # 추출해낸 데이터를 DataFrame 형식으로 변환한다.
        DTM_DataFrame = pd.DataFrame(DTM_Array, columns=feature_names)
            
        #if s_file_name in locals :
        # 불용어 열 제거
        #DTM_DataFrame.drop(stop_words_list, axis='columns', inplace=True)
            
        # 최종적으로 DTM을 csv 파일로 저장한다.
        DTM_DataFrame.to_csv('/home/ubuntu/web-codding/etc_file/fakenewsbody_DTM.csv', encoding='utf-8-sig')
        
        # DTM파일 가지고 동시출현단어.csv 파일 만들기
        dataset = pd.read_csv('/home/ubuntu/web-codding/etc_file/fakenewsbody_DTM.csv')
        
        column_list = dataset.columns[1:]
        word_length = len(column_list)
    
        count_dict={}
        
        for doc_number in tqdm(range(len(dataset)), desc='단어쌍 만들기 진행중') :
            tmp = dataset.loc[doc_number]
        for i, word1 in enumerate(column_list) :
            if tmp[word1] :
                for j in range(i+1, word_length) :
                    if tmp[column_list[j]] :
                        count_dict[column_list[i], column_list[j]] = count_dict.get((column_list[i]), 0) + max(tmp[word1], tmp[column_list[j]])
        
        count_list=[]
        for words in count_dict :
            count_list.append([words[0], words[1], count_dict[words]])
        
        df = pd.DataFrame(count_list, columns=["word1", "word2", "freq"])
        df = df.sort_values(by=['freq'], ascending=False) # 내림차순
        df = df.reset_index(drop=True)
            
        df.to_csv('/home/ubuntu/web-codding/etc_file/fakenews_networkx.csv', encoding='utf-8-sig')
        

# db에 빈 테이블 생성
def make_table() :
    # 빈 테이블 ex_apidata, ban_apidata, co_occurrence_word 생성
    conn = pymysql.connect(host='', user='', password='', db='', charset='utf8')

    sql_ban_apidata = "SET sql_mode=''; \
            CREATE TABLE ban_apidata ( PRDT_NM varchar(250) NOT NULL, \
            MUFC_NM VARCHAR(150) NOT NULL DEFAULT 'NOPE', \
            PRIMARY KEY (PRDT_NM, MUFC_NM) );"
    sql_coword = "SET sql_mode=''; \
            CREATE TABLE co_occurrence_word (word1 varchar(30) not null, \
            word2 varchar(30) not null, \
            count int(11), \
            primary key (word1, word2) );"
    sql_fakecoword = "SET sql_mode=''; \
            CREATE TABLE fake_co_occurrence_word (word1 varchar(30) not null, \
            word2 varchar(30) not null, \
            count int(11), \
            primary key (word1, word2) );"
    sql_ex_apidata="SET sql_mode=''; create table ex_apidata(PRDLST_NM varchar(250) not null, BSSH_NM varchar(150) not null default 'NOPE', primary key(PRDLST_NM,BSSH_NM));"

    try :
        cur = conn.cursor()
        cur.execute(sql_ex_apidata)
        print(' - api 생성완료')
        cur.execute(sql_ban_apidata)
        print(' - ban 생성완료')
        cur.execute(sql_coword)
        print(' - co 생성완료')
        cur.execute(sql_fakecoword)
        print(' - fco 생성완료')
    finally :
        cur.close()
        conn.close()

def check_table() :
    # mysql 조회 커리문 - '특정'데이터베이스에 '특정' 테이블이 있는지, 있으면 1, 없으면 0
    engine, conn = connect_sql()
    sql = "SELECT EXISTS ( SELECT 1 FROM Information_schema.tables WHERE table_schema='DB이름' AND table_name='ex_apidata') AS flag"
    df = pd.read_sql(sql, conn) # 쿼리문 결과를 데이터프레임으로 가져옴
    check = df.iloc[0,0] # check 타입은 정수
    return check

# api데이터 받아와서 db에 집어넣기

def put_Rapi() : 
    # ex_apidata
    # 데이터프레임에 api 전부 받아와 중복제거 후 ex_apidata에 저장
        
    try :
        total_df = pd.DataFrame(columns=['PRDLST_NM', 'BSSH_NM'])
        i=1
        j=1000
        while 1 :
            print("api받아오는중...")
            url = "http://openapi.foodsafetykorea.go.kr/api/개인키/I0030/json/"+str(i)+"/"+str(j)
            data = requests.get(url).json()
            if data['I0030']['RESULT']['MSG'] == "해당하는 데이터가 없습니다." :
                break
            else :
                body = [data['I0030']['row']]
                a = pd.json_normalize(data['I0030']['row'])
                # BSSH_NM : 업소_명, # PRDLST_NM : 품목_명
                info_df = a[['PRDLST_NM', 'BSSH_NM']]
                # print(info_df) # dataFrame
                info_df.columns = ['PRDLST_NM', 'BSSH_NM']
                total_df = pd.concat([total_df,info_df], axis=0, ignore_index=True)
                i+=1000
                j+=1000
        # 중복 데이터 제거
        total_df = total_df.drop_duplicates(['PRDLST_NM', 'BSSH_NM'], keep='first') #first:첫번째행남김, last:마지막행남김
        # 제품명 열 데이터 띄어쓰기 제거
        total_df['PRDLST_NM'] = total_df['PRDLST_NM'].str.replace(' ','')
        total_df['BSSH_NM'] = total_df['BSSH_NM'].str.replace(' ','')
    
    except :
        total_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/ex_apidata.csv')
        print("api - 다운로드 파일로 생성중")
        del total_df['Unnamed: 2']
    
    total_df.to_sql(name = 'ex_apidata', con=engine, if_exists='replace', index = False)
    return total_df


def put_Bapi() :
    # ban_apidata
    # 데이터프레임에 api 전부 받아와 중복제거 후 ban_apidata에 저장    
    try :
        bantotal_df = pd.DataFrame(columns=['PRDT_NM', 'MUFC_NM'])
        i=1
        j=1000
        while 1:
            url = "http://openapi.foodsafetykorea.go.kr/api/개인키/I2715/json/"+str(i)+"/"+str(j)
            data = requests.get(url).json()
            if data['I2715']['RESULT']['MSG'] == "해당하는 데이터가 없습니다." :
                break
            else :
                body = [data['I2715']['row']]
                a = pd.json_normalize(data['I2715']['row'])
                info_df = a[['PRDT_NM', 'MUFC_NM']]
                info_df.columns = ['PRDT_NM', 'MUFC_NM']
                bantotal_df = pd.concat([bantotal_df, info_df], axis=0, ignore_index=True)
                i+=1000
                j+=1000

        # 위해, 판매중지 식품
        fake_df = pd.DataFrame(columns=['PRDT_NM', 'MUFC_NM'])
        i=1
        j=1000
        while 1 :
            url = "http://openapi.foodsafetykorea.go.kr/api/개인키/I0490/json/"+str(i)+"/"+str(j)
            data = requests.get(url).json()
            if data['I0490']['RESULT']['MSG'] == "해당하는 데이터가 없습니다." :
                break
            else :
                body = [data['I0490']['row']]
                b = pd.json_normalize(data['I0490']['row'])
                fake_info_df = b[['PRDTNM', 'BSSHNM']]
                fake_info_df.columns = ['PRDT_NM', 'MUFC_NM']
                fake_df = pd.concat([fake_df, fake_info_df], axis=0, ignore_index=True)
                i+=1000
                j+=1000

        # 해외위해식품df 위해.판매중지df 합치기
        bantotal_df = pd.concat([bantotal_df, fake_df], axis=0, ignore_index=True)
        # 중복값 제거
        bantotal_df = bantotal_df.drop_duplicates(['PRDT_NM', 'MUFC_NM'], keep=False)
        bantotal_df = bantotal_df.reset_index(drop=True)
        # 제품명 열 데이터 띄어쓰기 제거
        bantotal_df['PRDT_NM'] = bantotal_df['PRDT_NM'].str.replace(' ','')
    except :
        print('ban - 다운로드 파일로 생성중..')
        bantotal_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/ban_apidata.csv')
        del bantotal_df['Unnamed: 2']

    bantotal_df.to_sql(name='ban_apidata', con=engine, if_exists='replace', index=False) # index=First&last 하면 오류남..
    return bantotal_df


def put_Cword() :
    # co_occurrence_word
    make_coword() # 동출빈.csv파일을 가져와야되니 우선 파일을 생성
    coword_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/_trustnews_networkx.csv')
    del coword_df['Unnamed: 0'] # 첫번째 쓸데없는 열 삭제
    coword_df.rename(columns={'freq':'count'}, inplace=True) # 열이름 변경
    coword_df.to_sql(name='co_occurrence_word', con=engine, if_exists='append', index=False)
    return coword_df
    

def put_FCword() : # fake_co_occurrence_word
    make_fcoword() # fake동출빈.csv파일을 가져와야되니 우선 파일을 생성
    fcoword_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/fakenews_networkx.csv')
    del fcoword_df['Unnamed: 0']
    fcoword_df.rename(columns={'freq':'count'}, inplace=True) # 열이름 변경
    fcoword_df.to_sql(name='fake_co_occurrence_word', con=engine, if_exists='append', index=False)
    return fcoword_df


def check_presence_data() : # 테이블에 데이터가 있는지 확인

    sql_api = "select count(*) from ex_apidata;"
    cf = pd.read_sql(sql_api, conn) # 쿼리문 결과를 데이터프레임으로 가져옴
    sqlapi = cf.iloc[0,0] # check 타입은 정수

    sql_ban = "select count(*) from ban_apidata;"
    cf = pd.read_sql(sql_ban, conn)
    sqlban = cf.iloc[0,0]

    sql_co = "select count(*) from co_occurrence_word"
    cf = pd.read_sql(sql_co, conn)
    sqlco = cf.iloc[0,0]

    sql_fco = "select count(*) from fake_co_occurrence_word"
    cf = pd.read_sql(sql_fco, conn)
    sqlfco = cf.iloc[0,0]

    return sqlapi, sqlban, sqlco, sqlfco


def take_Rapi() : # mysql에서 ex_apidata 테이블(데이터) 가져오기
    #total_df = pd.read_sql_table('ex_apidata', conn)
    total_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/ex_apidata.csv')
    return total_df
def take_Bapi() : # mysql에서 ban_apidata 테이블(데이터) 가져오기
    #bantotal_df1 = pd.read_sql_table('ban_apidata', conn)
    bantotal_df1 = pd.read_csv('/home/ubuntu/web-codding/etc_file/ban_apidata.csv')
    return bantotal_df1
def take_Cword() : # mysql에서 co_occurrence_word 테이블(데이터) 가져오기
    #coword_df = pd.read_sql_table('co_occurrence_word', conn)
    coword_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/_trustnews_networkx.csv')
    del coword_df['Unnamed: 0']
    return coword_df
def take_FCword() : # 계산된 fake동출빈파일.csv 데이터 가져오기
    #fcoword_df = pd.read_sql_table('fake_co_occurrence_word', conn)
    fcoword_df = pd.read_csv('/home/ubuntu/web-codding/etc_file/fakenews_networkx.csv')
    del fcoword_df['Unnamed: 0']
    #conn.close() # sql 마지막 연결
    return fcoword_df


# 허위과대광고문구 있는지 조회하는 함수
def fake_text_find(v) :
    # 허위과대문구 텍스트파일에서 리스트로 불러오기
    fake_ad_file = open("/home/ubuntu/web-codding/etc_file/fake_sentence.txt", 'r') # 로컬서버
    fake_ad_list=[]
    while True :
        fake_ad = fake_ad_file.readline()
        if not fake_ad :
            break
        fake_ad = fake_ad.strip()
        fake_ad_list.append(fake_ad)
    fake_ad_file.close()
    #fake_ad_list # type = list
    # 허위과대문구 본문에 있는지 조회
    fake_ad_count = 0
    #print('')
    for i in range(len(fake_ad_list)) :
        if fake_ad_list[i] in v :
            print("삑삑 : {0}".format(fake_ad_list[i]))
            fake_ad_count +=1 
    return fake_ad_count

def get_recommendations(article, cosine_sim, indices, data_df) :
    idx = indices[article] # 입력한 기사로 인덱스를 받아옴
    sim_scores = list(enumerate(cosine_sim[idx])) # 모든 기사에 대해 입력한 기사와의 유사도를 구함
    sim_scores = sorted(sim_scores, key=lambda x: x[1], reverse=True) # 유사도에 따라 기사들을 정렬
    sim_scores = sim_scores[1:6] # 가장 유사한 기사 5개를 받아옴
    movie_indices = [i[0] for i in sim_scores] # 가장 유사한 기사 5개의 인덱스를 받아옴
    return data_df['code'].iloc[movie_indices] # 가장 유사한 기사의 신뢰도(높or낮)를 리턴


# --------------------------------------------------------

# 시작페이지
@app.route('/') # /start_page 라고 주소창에 입력하면 start_page.html파일을 열어 시작창을 띄움
def start_page() :
    return render_template('start_page.html')

# 사용설명 페이지
@app.route('/manual', methods=['get'])
def manual_page() :
    return render_template('manual_page.html')

# 파일추가 페이지
@app.route('/add-file', methods=['get'])
def add_file_page() :
    return render_template('add_file_page.html')

# 결과 페이지 - file
@app.route('/result', methods=['GET', 'POST'])
def result_page() :
    if request.method =='POST' :
        try :
            f = request.files['file'] # 파일불러옴
            f.save(os.path.join("uploads", f.filename)) # 파일저장 # 로컬 서버
            # ------------------------------------------------------------ 
            url = '/home/ubuntu/web-codding/uploads/'+f.filename # 로컬 서버
        except :
            url = request.form['url']
        
        if url[-3:] == 'pdf' :
            v = convert_pdf_to_text(url)
        elif url[-3:] == 'hwp' :
            v = convert_hwp_to_text(url)
        elif url[-4:] == 'docx' :
            v = convert_docx_to_text(url)
        elif url[-3:] == 'txt' :
            v = convert_txt_to_text(url)
        elif url[:4] == 'http' :
            v = convert_link_to_text(url)
        
        body_text = v
        sentence = clean_sentence(v)
        # ------------------------------------------------------------ 1
        # Mecab으로 명사 추출
        mecab = Mecab()
        final = mecab.nouns(sentence) # mecab으로 명사 추출
        final # type(final) = list
        # ------------------------------------------------------------ 2
        # 추출된 명사들 중 빈도수가 3이상인 명사들만 보관
        word_list = pd.Series(final) # 시리즈로 변경
        result = word_list.value_counts() # 개수 구하기
        result_df = pd.DataFrame(result) # 데이터프레임으로 변경
        result_df.columns=['빈도수'] # 컬럼명 '빈도수'로 수정

        # 빈도수 3 이상인 단어만 데이터프레임으로 따로 보관
        freq = result_df[(result_df['빈도수'] >= 3)]
        count = pd.DataFrame(freq.index, columns = ['word']) # freq(df) count(df)로 새로 생성
        # ------------------------------------------------------------ 3

        # ------------------------------------------------------------ 4
        """
        checktable=0 # 테이블 유무 확인변수
        checktable = check_table()
        if checktable == 0 : # 빈 테이블이 없다면
            print('테이블 없으므로 생성')
            make_table()
            print(' - 테이블 생성 완료')
        else :
            print('*이미 테이블이 생성되어있으므로 테이블 생성 건너뜀*')
        # ------------------------------------------------------------ 5
        
        sqlapi, sqlban, sqlco, sqlfco = check_presence_data()

        if sqlapi == 0 :
            print('api테이블 데이터 없으므로 저장')
            total_df = put_Rapi()
            print(' - Rapi(등록)완료')
        else :
            print('*api테이블에 데이터가 있으므로 저장 건너뜀*')
            total_df = take_Rapi() # mysql에서 데이터가져와서 df에 저장

        if sqlban == 0 :
            print('ban테이블에 데이터 없으므로 저장')
            bantotal_df1 = put_Bapi()
            print('- Bapi(금지)완료')
        else :
            print('*ban테이블에 데이터가 있으므로 저장 건너뜀*')
            bantotal_df1 = take_Bapi() # mysql에서 데이터가져와서 df에 저장

        if sqlco == 0 :
            print('coword테이블에 데이터 없으므로 저장')
            df = put_Cword()
            print('- Cword(동출빈)완료')
        else :
            print('*coword테이블에 데이터가 있으므로 저장 건너뜀*')    
            df = take_Cword() # mysql에서 데이터가져와서 df에 저장

        if sqlfco == 0 :
            print('fcoword테이블에 데이터 없으므로 저장')
            df2 = put_FCword()
            print('- FCword(fake동출빈)완료')
        else :
            print('*fcoword테이블에 데이터가 있으므로 저장 건너뜀*')
            df2 = take_FCword() # mysql에서 데이터가져와서 df에 저장

        """
        total_df = take_Rapi()
        print('total_df완료')
        bantotal_df1 = take_Bapi()
        print('ban_df완료')
        df = take_Cword()
        print('co_df완료')
        df2 = take_FCword()
        print('fco_df완료')
        # ------------------------------------------------------------ 
        # 식약처 db랑 비교해서 일치하는 단어만 따로 리스트에 저장

        # 등록상품 
        total_prdlst = np.array(total_df['PRDLST_NM'].tolist())
        total_bssh = np.array(total_df['BSSH_NM'].tolist())

        # 위해상품
        ban_prdt = np.array(bantotal_df1['PRDT_NM'].tolist())
        ban_mufc = np.array(bantotal_df1['MUFC_NM'].tolist())

        # final1 중복값 제거
        final1 = final
        final1 = set(final1)
        sample = list(final1)

        # db에 일치하는 결과가 있는 단어만 따로 빼서 저장할 리스트
        sample_list_total = []
        sample_list_ban = []

        for i in range(len(sample)):
            if sample[i] in total_prdlst:
                if sample[i] not in sample_list_total:
                    sample_list_total.append(sample[i])
            if sample[i] in total_bssh:
                if sample[i] not in sample_list_total:
                    sample_list_total.append(sample[i])
            if sample[i] in ban_prdt:
                if sample[i] not in sample_list_ban:
                    sample_list_ban.append(sample[i])
            if sample[i] in ban_mufc:
                if sample[i] not in sample_list_ban:
                    sample_list_ban.append(sample[i])


        print('등록된 제품명 : {0}'.format(sample_list_total))
        print('등록된 위해제품명 : {0}'.format(sample_list_ban))

        # ------------------------------------------------------------ 7
        # 동시출현 쌍 만들기.. (한문서)
        co_word = {}
        for i in range(len(final)):
            for j in range(len(final)):
                a = final[i]
                b = final[j]
                if a == b: continue    # 둘이 같은 단어인 경우는 세지 않음
                if a > b: a, b = b, a  # a, b와 b, a 가 다르게 세어지는 것을 막기 위해 순서 고정
                co_word[a, b] = co_word.get((a, b), 0) + 1/2  # 실제로 센다

        co_word_list = list(co_word)          # co_word 사전에서 key부분 (단어쌍 부분)만 추출해서 리스트로 변환
        co_freq_list = list(co_word.values()) # co_word 사전에서 value부분 (나온 횟수)만 추출해서 리스트로 변환 

        # co_word_list 리스트를 df로 변환 
        co_word_df = pd.DataFrame(co_word_list, columns = ['word1', 'word2'])

        # df에 새 컬럼 'freq'를 추가하고, 값으로 co_freq_list (나온 횟수) 집어넣기
        co_word_df['freq'] = co_freq_list

        # 한 문서 내의 단어 쌍 df 
        #co_word_df

        # 신뢰도 높음 동출빈과 비교해 일치하는 것을 따로 저장
        df_merge = co_word_df.merge(df)
        #print("\n신뢰도 높음과 일치\n",df_merge)

        # 신뢰도 낮음 동출빈과 비교해 일치하는 것을 따로 저장
        df2_merge = co_word_df.merge(df2)
        #print("\n신뢰도 낮음과 일치\n",df2_merge)

        # -------------------------------------------------------- 8
        # 기사의 유사도를 구하는 부분
        
        #입력한 기사의 본문 불러오기
        tf_article = sentence

        # 기사자료.csv 불러오고 중복제거
        data = pd.read_csv('/home/ubuntu/web-codding/etc_file/article_data.csv', low_memory=False) # csv 읽어오기
        data_df = data[['article', 'code']] # data(df) 열이름 변경
        data_df.drop_duplicates() # 중복제거

        # 입력한 기사를 df에 추가
        data_insert = {'article':tf_article, 'code':'불명'} # 기사본문(1개)
        data_df = data_df.append(data_insert, ignore_index=True) # data_df(csv읽어온df)에 추가

        indices = pd.Series(data_df.index, index=data_df['article']) # 합친 기사본문 가지고 판다스 시리즈 생성

        tfidf = TfidfVectorizer()

        # title에 대해서 tf-idf 수행
        tfidf_matrix = tfidf.fit_transform(data_df['article'])
        print(tfidf_matrix.shape)

        cosine_sim = linear_kernel(tfidf_matrix, tfidf_matrix)

        # get_recommendations 함수의 결과를 받아올 df 생성
        get_df = pd.DataFrame(get_recommendations(tf_article, cosine_sim, indices, data_df))

        value_series = get_df['code'].value_counts()

        # 보통도 있을 수 있는데 보통은 계산에 포함X
        try :
            high = value_series['높음']
        except KeyError :
            high = 0
        try :
            low = value_series['낮음']
        except KeyError :
            low = 0
        print('high :', high,", low :", low)

        tf_num = (high - low)*6

        # -------------------------------------------------------- 9
        reliability = 0 # 신뢰도 넣을 변수
        fake_count = fake_text_find(v) # 함수 조회로 카운트 받기

        if len(sample_list_total) == 0 and len(sample_list_ban) == 0:
            print("신뢰도 측정 대상 외의 기사입니다.")
            exit()
        else:
            # 신뢰도 총 70점
            if sample_list_total and not sample_list_ban:
                reliability += 50
                print("등록점수 :",reliability)
            elif not sample_list_total and sample_list_ban:
                reliability += 20
                print("등록점수 :",reliability)
            else: # 둘 다 겹칠 때
                reliability += 20
                print("등록점수 :",reliability)
    
            # 동시출현단어 빈도수 비교 (최대 점수 30점)
            if len(df_merge.index) >= len(df2_merge.index):
                num = int(len(df2_merge.index) / len(df_merge.index) * 20)
                num = 20 - num
            elif len(df_merge.index) < len(df2_merge.index):
                num = int(len(df_merge.index) / len(df2_merge.index) * 20)

            # num이 30이 넘거나 음수인 경우 값 조정
            if num > 20:
                num = 20
            elif num < 0:
                num = 0       
            reliability += num
            print("동출빈점수 :", num)

            # tf-idf 유사도 점수적용
            reliability += tf_num
            print("tf-idf점수 :", tf_num)

            if fake_count != 0 : # 허위과대광고 문구가 본문에 있을경우
                reliability -= 1*fake_count
                print("허위과대광고문구가 {0}개 있으므로 {1}점 감소".format(fake_count, 1*fake_count))
            

            # result_page에 출력할 본문 정리
            body_text = body_text.replace("\n", '')
            body_text = body_text.replace("\x0c", '')
            list_bd = body_text.split(sep='. ') # ". "으로 한 문장씩 나눠서 리스트로 정리
            list_bd = [item + '.' for item in list_bd] # list로 나눠진 문장끝에 . 추가
            body_text = list_bd

    elif request.method =='GET' :
        body_text = "이무진 - 과제\n\n"
        reliability = 100
    return render_template('result_page.html', body_text=body_text, percentage=reliability, number=len(body_text), sample_list_total=sample_list_total, sample_list_ban=sample_list_ban) # html파일과 변수 return

@app.route('/evaluate', methods=['POST', 'GET'])
def evaluate_page() :
    if request.method == 'POST' :
        
        if request.form['options'] != None :
            option = request.form['options']
        else :
            option = "(기본값)"
        
        if request.form['text'] != None :
            text = request.form['text']
        else :
            text = "(기본값)"
        
        return render_template('evaluate_page.html', options=option, text=text)
    elif request.method == 'GET' :
        option = "(기본값)입력이 없어요~"
        text = "(기본값)입력이 없어요~"
        return render_template('evaluate_page.html', options=option, text=text)

# run the app
if __name__=="__main__" : 
    app.run(host='0.0.0.0', port=5000)
